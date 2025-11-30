import tkinter as tk
from tkinter import ttk, scrolledtext, filedialog, messagebox
import re
import sys
import os
import gc
import threading
from typing import List, Tuple, Dict, Optional

# 配置系统参数，适应大文件处理
sys.setrecursionlimit(100000)
os.environ["OMP_NUM_THREADS"] = "2"  # 限制多线程内存占用


class ThesisFormatChecker:
    def __init__(self, root: tk.Tk):
        self.root = root
        self.root.title("论文格式检查工具")
        self.root.geometry("1200x700")
        self.root.report_callback_exception = self._handle_exception  # 全局异常捕获

        # 论文专用格式规则（GB/T 7714等标准）
        self.format_rules = {
            "一级标题": {"中文字体": "黑体", "西文字体": "Times New Roman", "字号": "小二",
                         "段前距": 24, "段后距": 12, "行距": "固定值20", "行距值": 20,
                         "正则规则": r"^第[一二三四五六七八九十]+章 .+"},
            "二级标题": {"中文字体": "黑体", "西文字体": "Times New Roman", "字号": "小三",
                         "段前距": 20, "段后距": 10, "行距": "固定值20", "行距值": 20,
                         "正则规则": r"^[0-9]+\. .+"},
            "三级标题": {"中文字体": "黑体", "西文字体": "Times New Roman", "字号": "三号",
                         "段前距": 16, "段后距": 8, "行距": "固定值20", "行距值": 20,
                         "正则规则": r"^[0-9]+\.[0-9]+\. .+"},
            "摘要": {"中文字体": "宋体", "西文字体": "Times New Roman", "字号": "小三",
                     "段前距": 0, "段后距": 12, "行距": "固定值20", "行距值": 20,
                     "正则规则": r"^[摘|Abstract].*"},
            "关键词": {"中文字体": "宋体", "西文字体": "Times New Roman", "字号": "小四",
                       "段前距": 0, "段后距": 12, "行距": "固定值20", "行距值": 20,
                       "正则规则": r"^[关键|Key].*词.*"},
            "正文": {"中文字体": "宋体", "西文字体": "Times New Roman", "字号": "小四",
                     "段前距": 0, "段后距": 0, "行距": "固定值20", "行距值": 20,
                     "正则规则": r".+"},
            "参考文献": {"中文字体": "宋体", "西文字体": "Times New Roman", "字号": "小三",
                         "段前距": 12, "段后距": 0, "行距": "单倍行距", "行距值": 1.0,
                         "正则规则": r"^[0-9]+\. .+"},
            "表格标题": {"中文字体": "黑体", "西文字体": "Times New Roman", "字号": "五号",
                         "位置": "上方", "正则规则": r"^表\s*\d+.*"},
            "图片标题": {"中文字体": "黑体", "西文字体": "Times New Roman", "字号": "五号",
                         "位置": "下方", "正则规则": r"^图\s*\d+.*"}
        }

        # 行距映射表
        self.line_spacing_mapping = {
            "单倍行距": 1.0, "1.5倍行距": 1.5, "2倍行距": 2.0,
            "固定值20": 20, "固定值24": 24, "最小值": 1.0,
            "多倍行距1.2": 1.2, "多倍行距1.3": 1.3, "多倍行距1.4": 1.4
        }

        # 第三方库延迟加载
        self.docx = None
        self.pdfplumber = None

        # 大文件处理参数（核心优化：限制处理规模）
        self.batch_size = 50  # 分批处理的段落数量
        self.max_total_pages = 200  # 最大处理页数
        self.max_paragraphs = 300  # Word最大处理段落数
        self.max_pdf_pages = 20  # PDF最大处理页数（初期可设小些）

        # 创建界面
        self._create_widgets()

    def _handle_exception(self, exc_type, exc_value, exc_traceback):
        """全局异常处理，防止闪退并释放内存"""
        if issubclass(exc_type, KeyboardInterrupt):
            self.root.destroy()
            return
        error_msg = f"错误信息：{str(exc_value)[:150]}\n\n建议：尝试减小处理页数或更换文件"
        messagebox.showerror("处理错误", error_msg)
        gc.collect()  # 强制释放内存

    def _create_widgets(self) -> None:
        # 顶部标签页
        tab_control = ttk.Notebook(self.root)

        self.tab_ref = ttk.Frame(tab_control)
        tab_control.add(self.tab_ref, text="格式标准设置")

        self.tab_custom = ttk.Frame(tab_control)
        tab_control.add(self.tab_custom, text="自定义规则")

        self.tab_check = ttk.Frame(tab_control)
        tab_control.add(self.tab_check, text="论文检查")
        self._setup_check_tab()  # 优先初始化检查页

        self.tab_result = ttk.Frame(tab_control)
        tab_control.add(self.tab_result, text="检查结果")
        self._setup_result_tab()

        tab_control.pack(expand=1, fill="both")

        # 配置其他标签页
        self._setup_reference_tab()
        self._setup_custom_tab()

    # ------------------------------
    # 界面配置部分（含子线程启动按钮）
    # ------------------------------
    def _setup_check_tab(self) -> None:
        """论文检查界面（核心修改：子线程启动检查）"""
        input_frame = ttk.LabelFrame(self.tab_check, text="论文文件设置")
        input_frame.pack(fill="x", padx=5, pady=5)

        self.input_method = tk.StringVar(value="file")
        ttk.Radiobutton(input_frame, text="上传论文文件", variable=self.input_method,
                        value="file", command=self._toggle_input_method).pack(side="left", padx=10)
        ttk.Radiobutton(input_frame, text="粘贴文本片段", variable=self.input_method,
                        value="text", command=self._toggle_input_method).pack(side="left", padx=10)

        # 大文件处理选项
        options_frame = ttk.LabelFrame(self.tab_check, text="大文件处理选项")
        options_frame.pack(fill="x", padx=5, pady=5)

        ttk.Label(options_frame, text="处理页数限制:").pack(side="left", padx=5)
        self.page_limit = ttk.Combobox(
            options_frame,
            values=["全部", "50", "100", "200", "300"],
            width=6
        )
        self.page_limit.current(1)  # 默认50页
        self.page_limit.pack(side="left", padx=5)

        ttk.Label(options_frame, text="处理模式:").pack(side="left", padx=5)
        self.process_mode = ttk.Combobox(
            options_frame,
            values=["快速检查（仅文本）", "详细检查（含格式）"],
            width=15
        )
        self.process_mode.current(0)
        self.process_mode.pack(side="left", padx=5)

        # 文件选择区域
        self.file_frame = ttk.Frame(self.tab_check)
        self.file_label = ttk.Label(self.file_frame, text="未选择文件（支持大文件docx/pdf）")
        self.file_label.pack(side="left", padx=5)
        ttk.Button(self.file_frame, text="选择论文文件", command=self._select_file).pack(side="left", padx=5)
        self.file_path = tk.StringVar()

        # 文本输入区域
        self.text_input = scrolledtext.ScrolledText(self.tab_check, wrap=tk.WORD)

        # 进度条
        self.progress_frame = ttk.Frame(self.tab_check)
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(
            self.progress_frame,
            variable=self.progress_var,
            maximum=100
        )
        self.progress_label = ttk.Label(self.progress_frame, text="准备就绪")
        self.progress_bar.pack(side="left", fill="x", expand=True, padx=5)
        self.progress_label.pack(side="left", padx=5)
        self.progress_frame.pack(fill="x", padx=5, pady=5)

        # 核心修改：通过子线程启动检查，避免GUI阻塞
        ttk.Button(
            self.tab_check,
            text="开始论文格式检查",
            command=self._start_check_thread,  # 启动子线程
            width=20
        ).pack(pady=10)

        self._toggle_input_method()

    def _start_check_thread(self):
        """新增：用子线程执行检查逻辑，防止GUI卡顿闪退"""
        threading.Thread(
            target=self._start_check,
            daemon=True  # 主线程退出时子线程也退出
        ).start()

    # ------------------------------
    # 其他界面配置（保持不变）
    # ------------------------------
    def _setup_reference_tab(self) -> None:
        ttk.Label(self.tab_ref, text="格式类别:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.format_type = ttk.Combobox(self.tab_ref, values=list(self.format_rules.keys()), width=18)
        self.format_type.current(0)
        self.format_type.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        self.format_type.bind("<<ComboboxSelected>>", self._load_format_values)

        props_frame = ttk.LabelFrame(self.tab_ref, text="格式属性（符合GB/T 7714等标准）")
        props_frame.grid(row=1, column=0, columnspan=4, padx=5, pady=5, sticky="nsew")

        ttk.Label(props_frame, text="中文字体:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.chinese_font = ttk.Entry(props_frame, width=15)
        self.chinese_font.grid(row=0, column=1, padx=5, pady=5, sticky="w")

        ttk.Label(props_frame, text="西文字体:").grid(row=0, column=2, padx=5, pady=5, sticky="w")
        self.english_font = ttk.Entry(props_frame, width=15)
        self.english_font.grid(row=0, column=3, padx=5, pady=5, sticky="w")

        ttk.Label(props_frame, text="字号(磅):").grid(row=1, column=0, padx=5, pady=5, sticky="w")
        self.font_size = ttk.Entry(props_frame, width=10)
        self.font_size.grid(row=1, column=1, padx=5, pady=5, sticky="w")

        ttk.Label(props_frame, text="段前距(磅):").grid(row=1, column=2, padx=5, pady=5, sticky="w")
        self.space_before = ttk.Entry(props_frame, width=10)
        self.space_before.grid(row=1, column=3, padx=5, pady=5, sticky="w")

        ttk.Label(props_frame, text="段后距(磅):").grid(row=2, column=0, padx=5, pady=5, sticky="w")
        self.space_after = ttk.Entry(props_frame, width=10)
        self.space_after.grid(row=2, column=1, padx=5, pady=5, sticky="w")

        ttk.Label(props_frame, text="行距:").grid(row=2, column=2, padx=5, pady=5, sticky="w")
        self.line_spacing = ttk.Combobox(
            props_frame,
            values=["单倍行距", "1.5倍行距", "2倍行距", "固定值20", "固定值24", "多倍行距1.2"],
            width=12
        )
        self.line_spacing.grid(row=2, column=3, padx=5, pady=5, sticky="w")

        ttk.Label(props_frame, text="标题位置:").grid(row=3, column=0, padx=5, pady=5, sticky="w")
        self.caption_pos = ttk.Combobox(props_frame, values=["上方", "下方"], width=8)
        self.caption_pos.grid(row=3, column=1, padx=5, pady=5, sticky="w")

        ttk.Label(props_frame, text="匹配规则(正则):").grid(row=4, column=0, padx=5, pady=5, sticky="w")
        self.regex_rule = ttk.Entry(props_frame, width=60)
        self.regex_rule.grid(row=4, column=1, columnspan=3, padx=5, pady=5, sticky="we")
        ttk.Label(props_frame, text="示例：^第[一二]+章 绪论", font=("Arial", 8)).grid(
            row=4, column=4, padx=5, pady=5, sticky="w")

        ttk.Button(self.tab_ref, text="保存格式设置", command=self._save_format_settings).grid(
            row=5, column=0, padx=5, pady=10)
        ttk.Button(self.tab_ref, text="加载论文标准格式", command=self._load_default_formats).grid(
            row=5, column=1, padx=5, pady=10)

        self.tab_ref.grid_columnconfigure(3, weight=1)
        self.tab_ref.grid_rowconfigure(1, weight=1)
        props_frame.grid_columnconfigure(4, weight=1)
        self._load_format_values(None)

    def _setup_custom_tab(self) -> None:
        add_frame = ttk.LabelFrame(self.tab_custom, text="新增论文格式规则")
        add_frame.grid(row=0, column=0, columnspan=2, padx=5, pady=5, sticky="we")

        ttk.Label(add_frame, text="规则名称:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.new_rule_name = ttk.Entry(add_frame, width=20)
        self.new_rule_name.grid(row=0, column=1, padx=5, pady=5, sticky="w")

        ttk.Label(add_frame, text="匹配正则:").grid(row=0, column=2, padx=5, pady=5, sticky="w")
        self.new_rule_regex = ttk.Entry(add_frame, width=50)
        self.new_rule_regex.grid(row=0, column=3, padx=5, pady=5, sticky="we")

        ttk.Button(add_frame, text="添加规则", command=self._add_custom_rule).grid(row=0, column=4, padx=10, pady=5)

        list_frame = ttk.LabelFrame(self.tab_custom, text="已存在规则（论文常用）")
        list_frame.grid(row=1, column=0, columnspan=2, padx=5, pady=5, sticky="nsew")

        self.rule_listbox = tk.Listbox(list_frame, width=100, height=15)
        self.rule_listbox.grid(row=0, column=0, padx=5, pady=5, sticky="nsew")
        scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=self.rule_listbox.yview)
        scrollbar.grid(row=0, column=1, sticky="ns")
        self.rule_listbox.config(yscrollcommand=scrollbar.set)

        ttk.Button(self.tab_custom, text="删除选中规则", command=self._delete_custom_rule).grid(
            row=2, column=0, padx=5, pady=10, sticky="w")
        ttk.Button(self.tab_custom, text="刷新规则列表", command=self._refresh_rule_list).grid(
            row=2, column=1, padx=5, pady=10, sticky="w")

        self.tab_custom.grid_columnconfigure(0, weight=1)
        self.tab_custom.grid_rowconfigure(1, weight=1)
        add_frame.grid_columnconfigure(3, weight=1)
        list_frame.grid_columnconfigure(0, weight=1)
        self._refresh_rule_list()

    def _setup_result_tab(self) -> None:
        filter_frame = ttk.LabelFrame(self.tab_result, text="结果过滤")
        filter_frame.pack(fill="x", padx=5, pady=5)

        self.filter_type = ttk.Combobox(
            filter_frame,
            values=["全部问题", "字体错误", "字号错误", "行距错误", "未匹配规则"],
            width=15
        )
        self.filter_type.current(0)
        self.filter_type.pack(side="left", padx=10)
        ttk.Button(filter_frame, text="应用过滤", command=self._filter_results).pack(side="left", padx=5)

        self.result_display = scrolledtext.ScrolledText(
            self.tab_result,
            wrap=tk.WORD,
            state="disabled",
            width=120,
            height=30
        )
        self.result_display.pack(fill="both", expand=True, padx=5, pady=5)

        self.stats_label = ttk.Label(self.tab_result, text="问题统计：共0个问题")
        self.stats_label.pack(side="left", padx=5, pady=5)

        ttk.Button(self.tab_result, text="导出详细报告", command=self._export_results).pack(side="right", padx=5,
                                                                                            pady=5)
        self.raw_issues = []

    # ------------------------------
    # 格式规则管理（保持不变）
    # ------------------------------
    def _load_format_values(self, event) -> None:
        fmt_type = self.format_type.get()
        fmt = self.format_rules.get(fmt_type, {})

        self.chinese_font.delete(0, tk.END)
        self.chinese_font.insert(0, fmt.get("中文字体", ""))

        self.english_font.delete(0, tk.END)
        self.english_font.insert(0, fmt.get("西文字体", ""))

        self.font_size.delete(0, tk.END)
        self.font_size.insert(0, fmt.get("字号", ""))

        self.space_before.delete(0, tk.END)
        self.space_before.insert(0, fmt.get("段前距", ""))

        self.space_after.delete(0, tk.END)
        self.space_after.insert(0, fmt.get("段后距", ""))

        self.line_spacing.set(fmt.get("行距", "1.5倍行距"))
        self.caption_pos.set(fmt.get("位置", ""))

        self.regex_rule.delete(0, tk.END)
        self.regex_rule.insert(0, fmt.get("正则规则", ""))

    def _parse_line_spacing(self, spacing_text: str) -> float:
        try:
            return float(spacing_text)
        except ValueError:
            return self.line_spacing_mapping.get(spacing_text, 1.5)

    def _save_format_settings(self) -> None:
        fmt_type = self.format_type.get()
        try:
            font_size = int(self.font_size.get()) if self.font_size.get().isdigit() else 12
            space_before = int(self.space_before.get()) if self.space_before.get().isdigit() else 0
            space_after = int(self.space_after.get()) if self.space_after.get().isdigit() else 0
            line_spacing_text = self.line_spacing.get()
            line_spacing_value = self._parse_line_spacing(line_spacing_text)

            self.format_rules[fmt_type] = {
                "中文字体": self.chinese_font.get() or "宋体",
                "西文字体": self.english_font.get() or "Times New Roman",
                "字号": font_size,
                "段前距": space_before,
                "段后距": space_after,
                "行距": line_spacing_text,
                "行距值": line_spacing_value,
                "位置": self.caption_pos.get(),
                "正则规则": self.regex_rule.get().strip() or r".+"
            }
            self._show_message(f"{fmt_type} 格式已保存")
            self._refresh_rule_list()
        except Exception as e:
            self._show_message(f"保存失败：{str(e)}")

    def _load_default_formats(self) -> None:
        self.format_rules = {
            "一级标题": {"中文字体": "黑体", "西文字体": "Times New Roman", "字号": 24,
                         "段前距": 24, "段后距": 12, "行距": "1.5倍行距", "行距值": 1.5,
                         "正则规则": r"^第[一二三四五六七八九十]+章 .+"},
            "二级标题": {"中文字体": "黑体", "西文字体": "Times New Roman", "字号": 20,
                         "段前距": 20, "段后距": 10, "行距": "1.5倍行距", "行距值": 1.5,
                         "正则规则": r"^[0-9]+\. .+"},
            "三级标题": {"中文字体": "黑体", "西文字体": "Times New Roman", "字号": 16,
                         "段前距": 16, "段后距": 8, "行距": "1.5倍行距", "行距值": 1.5,
                         "正则规则": r"^[0-9]+\.[0-9]+\. .+"},
            "摘要": {"中文字体": "宋体", "西文字体": "Times New Roman", "字号": 12,
                     "段前距": 0, "段后距": 12, "行距": "1.5倍行距", "行距值": 1.5,
                     "正则规则": r"^[摘|Abstract].*"},
            "关键词": {"中文字体": "宋体", "西文字体": "Times New Roman", "字号": 12,
                       "段前距": 0, "段后距": 12, "行距": "1.5倍行距", "行距值": 1.5,
                       "正则规则": r"^[关键|Key].*词.*"},
            "正文": {"中文字体": "宋体", "西文字体": "Times New Roman", "字号": "小四",
                     "段前距": 0, "段后距": 0, "行距": "固定值", "行距值": 20,
                     "正则规则": r".+"},
            "参考文献": {"中文字体": "宋体", "西文字体": "Times New Roman", "字号": 10.5,
                         "段前距": 12, "段后距": 0, "行距": "单倍行距", "行距值": 1.0,
                         "正则规则": r"^[0-9]+\. .+"},
            "表格标题": {"中文字体": "宋体", "西文字体": "Times New Roman", "字号": 10.5,
                         "位置": "上方", "正则规则": r"^表\s*\d+.*"},
            "图片标题": {"中文字体": "宋体", "西文字体": "Times New Roman", "字号": 10.5,
                         "位置": "下方", "正则规则": r"^图\s*\d+.*"}
        }
        self.format_type["values"] = list(self.format_rules.keys())
        self._load_format_values(None)
        self._show_message("已加载论文标准格式（符合GB/T 7714）")
        self._refresh_rule_list()

    def _add_custom_rule(self) -> None:
        rule_name = self.new_rule_name.get().strip()
        rule_regex = self.new_rule_regex.get().strip()

        if not rule_name:
            messagebox.showwarning("警告", "规则名称不能为空！")
            return
        if rule_name in self.format_rules:
            messagebox.showwarning("警告", f"已存在'{rule_name}'规则，请更换名称！")
            return
        if not rule_regex:
            messagebox.showwarning("警告", "正则匹配规则不能为空！")
            return

        self.format_rules[rule_name] = {
            "中文字体": "宋体", "西文字体": "Times New Roman", "字号": "小四",
            "段前距": 0, "段后距": 0, "行距": "固定值", "行距值": 20,
            "位置": "上方", "正则规则": rule_regex
        }

        self.format_type["values"] = list(self.format_rules.keys())
        self.new_rule_name.delete(0, tk.END)
        self.new_rule_regex.delete(0, tk.END)

        messagebox.showinfo("成功", f"新增'{rule_name}'规则成功！")
        self._refresh_rule_list()

    def _refresh_rule_list(self) -> None:
        self.rule_listbox.delete(0, tk.END)
        for rule_name, rule_info in self.format_rules.items():
            regex = rule_info.get("正则规则", "无匹配规则")
            display_text = f"{rule_name}：{regex[:50]}{'...' if len(regex) > 50 else ''}"
            self.rule_listbox.insert(tk.END, display_text)

    def _delete_custom_rule(self) -> None:
        selected_indices = self.rule_listbox.curselection()
        if not selected_indices:
            messagebox.showwarning("警告", "请先选中要删除的规则！")
            return

        selected_text = self.rule_listbox.get(selected_indices[0])
        rule_name = selected_text.split("：")[0].strip()

        system_rules = ["一级标题", "二级标题", "三级标题", "正文", "参考文献"]
        if rule_name in system_rules:
            messagebox.showwarning("保护提示", f"'{rule_name}'是系统规则，不能删除！")
            return

        if messagebox.askyesno("确认", f"确定要删除'{rule_name}'规则吗？"):
            if rule_name in self.format_rules:
                del self.format_rules[rule_name]
                self.format_type["values"] = list(self.format_rules.keys())
                self._refresh_rule_list()
                messagebox.showinfo("成功", f"已删除'{rule_name}'规则")

    # ------------------------------
    # 文件处理部分（核心优化：限制读取规模）
    # ------------------------------
    def _toggle_input_method(self) -> None:
        if self.input_method.get() == "text":
            self.text_input.pack(fill="both", expand=True, padx=5, pady=5)
            self.file_frame.pack_forget()
        else:
            self.text_input.pack_forget()
            self.file_frame.pack(fill="x", padx=5, pady=5)

    def _select_file(self) -> None:
        file_path = filedialog.askopenfilename(
            filetypes=[("论文文件", "*.docx *.pdf"),
                       ("Word文件", "*.docx"),
                       ("PDF文件", "*.pdf")]
        )
        if file_path:
            file_size = os.path.getsize(file_path) / (1024 * 1024)
            if file_size > 200:
                messagebox.showwarning("文件过大", "建议将文件拆分后检查（单文件限200MB）")
                return
            self.file_path.set(file_path)
            file_name = os.path.basename(file_path)
            self.file_label.config(text=f"已选择：{file_name}（{file_size:.1f}MB）")

    def _update_progress(self, value: float, message: str) -> None:
        self.progress_var.set(value)
        self.progress_label.config(text=message)
        self.root.update_idletasks()

    def _start_check(self) -> None:
        """文件检查主逻辑（在子线程中执行）"""
        self._update_progress(0, "准备处理...")
        self.raw_issues = []
        content_data = []
        file_type = ""

        try:
            page_limit_text = self.page_limit.get()
            self.max_total_pages = int(page_limit_text) if page_limit_text != "全部" else 1000

            if self.input_method.get() == "text":
                text = self.text_input.get("1.0", tk.END)
                content_data = [(line.strip(), {}) for line in text.split("\n") if line.strip()]
                file_type = "text"
                self._update_progress(30, "正在分析文本片段...")

            else:
                file_path = self.file_path.get()
                if not file_path:
                    self._show_message("请先选择论文文件")
                    self._update_progress(0, "准备就绪")
                    return

                if file_path.endswith(".docx"):
                    if not self.docx:
                        import docx
                        self.docx = docx
                    doc = self.docx.Document(file_path)
                    total_paragraphs = len(doc.paragraphs)
                    # 核心优化：限制最大处理段落数（取设置值和实际段落数的最小值）
                    process_paragraphs = min(total_paragraphs, self.max_paragraphs)
                    self._update_progress(10, f"准备处理Word文件（共{process_paragraphs}段）")

                    # 分批读取段落
                    for i in range(0, process_paragraphs, self.batch_size):
                        batch_paragraphs = doc.paragraphs[i:i + self.batch_size]
                        batch_data = self._process_docx_batch(batch_paragraphs)
                        content_data.extend(batch_data)
                        progress = 10 + (i / process_paragraphs) * 50
                        self._update_progress(progress,
                                              f"已处理 {i // self.batch_size + 1}/{(process_paragraphs // self.batch_size) + 1} 批")
                        gc.collect()  # 每批处理后释放内存
                    file_type = "docx"

                elif file_path.endswith(".pdf"):
                    if not self.pdfplumber:
                        import pdfplumber
                        self.pdfplumber = pdfplumber
                    # 核心优化：限制最大处理页数
                    with self.pdfplumber.open(file_path) as pdf:
                        total_pages = min(len(pdf.pages), self.max_pdf_pages)
                    self._update_progress(10, f"准备处理PDF文件（共{total_pages}页）")

                    # 分批读取页面
                    with self.pdfplumber.open(file_path) as pdf:
                        for i in range(0, total_pages, 5):
                            batch_pages = pdf.pages[i:i + 5]
                            batch_data = self._process_pdf_batch(batch_pages)
                            content_data.extend(batch_data)
                            progress = 10 + (i / total_pages) * 50
                            self._update_progress(progress, f"已处理 {i + 5}/{total_pages} 页")
                            gc.collect()  # 每批处理后释放内存
                    file_type = "pdf"

            self._update_progress(70, "正在检查格式问题...")
            if content_data:
                issues = self._check_format(content_data, file_type)
                self.raw_issues = issues
                self._display_results(issues)
                self.stats_label.config(text=f"问题统计：共{len(issues) - 2}个问题（总检查{len(content_data)}行）")
            else:
                self._show_message("未获取到文档内容")

            self._update_progress(100, "处理完成")
            self.root.after(3000, lambda: self._update_progress(0, "准备就绪"))

        except Exception as e:
            self._update_progress(0, "处理失败")
            raise e

    def _process_docx_batch(self, paragraphs: List) -> List[Tuple[str, Dict]]:
        """处理Word批次段落（限制单段长度）"""
        batch_data = []
        for para in paragraphs:
            if para.text.strip() and len(para.text) <= 2000:  # 过滤超长段落
                fmt_info = {}
                if self.process_mode.get() == "详细检查（含格式）" and para.runs:
                    chinese_fonts = set()
                    english_fonts = set()

                    for run in para.runs[:20]:  # 限制文本块处理数量
                        if run.font and run.font.name:
                            has_chinese = any('\u4e00' <= c <= '\u9fff' for c in run.text)
                            has_english = any(c.isascii() and c.isalpha() for c in run.text)

                            if has_chinese:
                                chinese_fonts.add(run.font.name)
                            if has_english:
                                english_fonts.add(run.font.name)

                    if chinese_fonts:
                        fmt_info['chinese_font'] = chinese_fonts.pop()
                    if english_fonts:
                        fmt_info['english_font'] = english_fonts.pop()

                    if para.runs[0].font.size:
                        fmt_info['size'] = para.runs[0].font.size.pt

                    fmt_info[
                        'space_before'] = para.paragraph_format.space_before.pt if para.paragraph_format.space_before else 0
                    fmt_info[
                        'space_after'] = para.paragraph_format.space_after.pt if para.paragraph_format.space_after else 0
                    fmt_info[
                        'line_spacing'] = para.paragraph_format.line_spacing if para.paragraph_format.line_spacing else 0

                batch_data.append((para.text[:1000], fmt_info))
        return batch_data

    def _process_pdf_batch(self, pages: List) -> List[Tuple[str, Dict]]:
        """处理PDF批次页面（限制每页文字数量）"""
        batch_data = []
        for page in pages:
            try:
                words = page.extract_words()[:200]  # 限制每页处理的文字数量
                if not words:
                    continue

                lines = []
                current_line = []
                current_y = None
                for word in words:
                    word_y = round(word['top'], 1)
                    if current_y is None:
                        current_y = word_y
                        current_line.append(word)
                    elif abs(word_y - current_y) < 2:
                        current_line.append(word)
                    else:
                        lines.append(current_line)
                        current_line = [word]
                        current_y = word_y
                if current_line:
                    lines.append(current_line)

                for line in lines:
                    if not line:
                        continue
                    line_text = ' '.join([word['text'] for word in line])
                    if not line_text.strip():
                        continue

                    fmt_info = {}
                    if self.process_mode.get() == "详细检查（含格式）":
                        chinese_fonts = set()
                        english_fonts = set()
                        current_size = None

                        for word in line[:30]:  # 限制每行处理的文字数量
                            font = word.get('fontname', '').split('+')[-1]
                            current_size = word.get('size', 0)

                            has_chinese = any('\u4e00' <= c <= '\u9fff' for c in word['text'])
                            has_english = any(c.isascii() and c.isalpha() for c in word['text'])

                            if has_chinese:
                                chinese_fonts.add(font)
                            if has_english:
                                english_fonts.add(font)

                        if chinese_fonts:
                            fmt_info['chinese_font'] = chinese_fonts.pop()
                        if english_fonts:
                            fmt_info['english_font'] = english_fonts.pop()
                        fmt_info['size'] = current_size

                    batch_data.append((line_text[:1000], fmt_info))

            except Exception as e:
                self._show_message(f"处理PDF页面出错：{str(e)[:30]}")
                continue
        return batch_data

    # ------------------------------
    # 格式检查与结果展示（保持不变）
    # ------------------------------
    def _check_format(self, content_data: List[Tuple[str, Dict]], file_type: str) -> List[str]:
        issues = []
        pdf_font_mapping = {
            "SimHei": "黑体", "SimSun": "宋体", "NSimSun": "新宋体",
            "Microsoft YaHei": "微软雅黑", "TimesNewRoman": "Times New Roman",
            "ArialMT": "Arial", "SimSun-ExtB": "宋体"
        }

        compiled_rules = {
            name: (re.compile(info["正则规则"]), info)
            for name, info in self.format_rules.items() if info.get("正则规则")
        }

        for i, (line, fmt_info) in enumerate(content_data, 1):
            if not line or len(line) > 1000:
                continue

            matched = False
            for rule_name, (pattern, rule_info) in compiled_rules.items():
                if pattern.match(line):
                    matched = True
                    if self.process_mode.get() == "详细检查（含格式）":
                        if "中文字体" in rule_info and rule_info["中文字体"]:
                            expected_cn_font = rule_info["中文字体"]
                            actual_cn_font = fmt_info.get('chinese_font', '未知')

                            if file_type == "pdf" and actual_cn_font in pdf_font_mapping:
                                actual_cn_font = pdf_font_mapping[actual_cn_font]

                            if actual_cn_font != expected_cn_font and actual_cn_font != '未知':
                                issues.append(
                                    f"行 {i}: {rule_name} 中文字体错误 - 应为'{expected_cn_font}'，实际为'{actual_cn_font}'")

                        if "西文字体" in rule_info and rule_info["西文字体"]:
                            expected_en_font = rule_info["西文字体"]
                            actual_en_font = fmt_info.get('english_font', '未知')

                            if file_type == "pdf" and actual_en_font in pdf_font_mapping:
                                actual_en_font = pdf_font_mapping[actual_en_font]

                            if actual_en_font != expected_en_font and actual_en_font != '未知':
                                issues.append(
                                    f"行 {i}: {rule_name} 西文字体错误 - 应为'{expected_en_font}'，实际为'{actual_en_font}'")

                        if "字号" in rule_info:
                            expected_size = rule_info["字号"]
                            actual_size = fmt_info.get('size', 0)

                            if actual_size > 0:
                                tolerance = 0.5 if expected_size <= 12 else 1.0
                                if not (expected_size - tolerance <= actual_size <= expected_size + tolerance):
                                    issues.append(
                                        f"行 {i}: {rule_name} 字号错误 - 应为'{expected_size}'，实际为'{actual_size:.1f}'")

                        if file_type == "docx":
                            if "段前距" in rule_info:
                                expected_before = rule_info["段前距"]
                                actual_before = fmt_info.get('space_before', 0)
                                if actual_before > 0 and not (
                                        expected_before - 2 <= actual_before <= expected_before + 2):
                                    issues.append(
                                        f"行 {i}: {rule_name} 段前距错误 - 应为'{expected_before}'，实际为'{actual_before:.1f}'")

                            if "段后距" in rule_info:
                                expected_after = rule_info["段后距"]
                                actual_after = fmt_info.get('space_after', 0)
                                if actual_after > 0 and not (expected_after - 2 <= actual_after <= expected_after + 2):
                                    issues.append(
                                        f"行 {i}: {rule_name} 段后距错误 - 应为'{expected_after}'，实际为'{actual_after:.1f}'")

                        if file_type == "docx" and "行距值" in rule_info:
                            expected_spacing = rule_info["行距值"]
                            actual_spacing = fmt_info.get('line_spacing', 0)

                            if actual_spacing > 0:
                                if expected_spacing <= 2.0:
                                    if not (expected_spacing - 0.2 <= actual_spacing <= expected_spacing + 0.2):
                                        issues.append(
                                            f"行 {i}: {rule_name} 行距错误 - 应为'{rule_info['行距']}'，实际为'{actual_spacing:.1f}倍'")
                                else:
                                    if not (expected_spacing - 2 <= actual_spacing <= expected_spacing + 2):
                                        issues.append(
                                            f"行 {i}: {rule_name} 行距错误 - 应为'{rule_info['行距']}'，实际为'{actual_spacing:.1f}'")

                        if "位置" in rule_info and rule_info["位置"]:
                            expected_pos = rule_info["位置"]
                            issues.append(f"行 {i}: {rule_name} 标题应位于{expected_pos}（需手动确认）")

                    break

            if not matched and line:
                issues.append(f"行 {i}: 未匹配任何格式规则 - '{line[:50]}{'...' if len(line) > 50 else ''}'")

        issues.append(f"\n共检测到 {len(content_data)} 行内容（处理模式：{self.process_mode.get()}）")
        if "正文" in self.format_rules:
            issues.append(
                f"正文标准格式: 中文字体{self.format_rules['正文']['中文字体']}，西文字体{self.format_rules['正文']['西文字体']}，"
                f"{self.format_rules['正文']['字号']}号字，行距{self.format_rules['正文']['行距']}")

        return issues[:500]

    def _display_results(self, issues: List[str]) -> None:
        self.result_display.config(state="normal")
        self.result_display.delete("1.0", tk.END)

        if issues:
            self.result_display.insert(tk.END, "论文格式检查结果：\n\n")
            for i in range(0, len(issues), 50):
                batch = issues[i:i + 50]
                for issue in batch:
                    self.result_display.insert(tk.END, f"- {issue}\n")
                self.root.update_idletasks()
        else:
            self.result_display.insert(tk.END, "未发现格式问题")

        self.result_display.config(state="disabled")

    def _filter_results(self) -> None:
        if not self.raw_issues:
            return

        filter_type = self.filter_type.get()
        filtered = []

        for issue in self.raw_issues:
            if filter_type == "全部问题":
                filtered.append(issue)
            elif filter_type == "字体错误" and ("中文字体错误" in issue or "西文字体错误" in issue):
                filtered.append(issue)
            elif filter_type == "字号错误" and "字号错误" in issue:
                filtered.append(issue)
            elif filter_type == "行距错误" and "行距错误" in issue:
                filtered.append(issue)
            elif filter_type == "未匹配规则" and "未匹配任何格式规则" in issue:
                filtered.append(issue)

        if len(self.raw_issues) >= 2:
            stats_lines = self.raw_issues[-2:]
            filtered = filtered[:-2] + stats_lines

        self._display_results(filtered)
        self.stats_label.config(text=f"过滤后：共{len(filtered) - 2}个问题（原{len(self.raw_issues) - 2}个）")

    def _export_results(self) -> None:
        if not self.raw_issues:
            self._show_message("没有可导出的结果")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("文本报告", "*.txt")]
        )

        if file_path:
            try:
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write("论文格式检查详细报告\n")
                    f.write("=" * 50 + "\n\n")
                    for i in range(0, len(self.raw_issues), 100):
                        batch = self.raw_issues[i:i + 100]
                        for issue in batch:
                            f.write(f"- {issue}\n")
                self._show_message(f"报告已导出至: {file_path}")
            except Exception as e:
                self._show_message(f"导出失败: {str(e)}")

    def _show_message(self, message: str) -> None:
        for widget in self.root.winfo_children():
            if isinstance(widget, ttk.Label) and widget["relief"] == tk.SUNKEN:
                widget.destroy()
        status_bar = ttk.Label(
            self.root,
            text=message,
            relief=tk.SUNKEN,
            anchor=tk.W
        )
        status_bar.pack(side=tk.BOTTOM, fill=tk.X)
        self.root.after(3000, status_bar.destroy)


if __name__ == "__main__":
    root = tk.Tk()
    app = ThesisFormatChecker(root)
    root.mainloop()