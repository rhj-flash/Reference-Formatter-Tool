# reference_processor.py

import re
import html
import os
from io import StringIO

from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# Pygments components for stable HTML formatting and segmentation
from pygments.formatters import HtmlFormatter
from pygments.lexer import RegexLexer
from pygments.token import Text

# 添加 Word 文档处理相关的导入
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from htmldocx import HtmlToDocx

    DOCX_AVAILABLE = True
    print("DEBUG: python-docx and htmldocx imported successfully")
except ImportError as e:
    print(f"WARNING: python-docx or htmldocx not available: {e}")
    DOCX_AVAILABLE = False

# 定义一个字典，用于进行全角到半角的转换
CHAR_MAPPING = {
    # 数字
    '０': '0', '１': '1', '２': '2', '３': '3', '四': '4',
    '５': '5', '６': '6', '７': '7', '８': '8', '９': '9',
    # 括号和方括号
    '（': '(', '）': ')', '【': '[', '】': ']',
    # 分隔符
    '—': '-', '–': '-', '―': '-', '．': '.', '：': ':', '；': ';',
    # 其他常见全角符号
    '，': ',', '。': '.', '？': '?', '！': '!', '＠': '@', '＃': '#',
    '＄': '$', '％': '%', '＾': '^', '＆': '&', '＊': '*', '＋': '+',
    '＝': '=', '～': '~', '　': ' '
}


# --- Pygments 自定义 Lexer (保持原样) ---

class ReferenceBlockLexer(RegexLexer):
    """
    一个简单的 Pygments Lexer，用于将每一行文本识别为一个块。
    我们只利用 Pygments 的分块能力来配合 Formatter。
    """
    name = 'RefBlock'
    aliases = ['refblock']
    filenames = ['*.ref']

    tokens = {
        'root': [
            # 匹配一整行，并将其标记为 Text.Line 类型
            (r'.*\n', Text.Line),
            (r'.*$', Text.Line),  # 匹配最后一行没有换行符的情况
        ],
    }


# --- Pygments 自定义 Formatter (保持原样) ---

# In reference_processor.py

class ReferenceBlockFormatter(HtmlFormatter):
    """
    自定义 Pygments HTML Formatter，用于显示格式化后的文献分割预览。
    支持编号格式和分割线显示。
    """
    # 启用全行着色
    full_lines = True

    def __init__(self, **options):
        super().__init__(style='default', **options)
        # 定义文献块样式
        self.BLOCK_STYLES = [
            {
                'bg_color': '#f0f8ff',  # 浅蓝色
                'border_color': '#b0d0e0',
                'header_bg': '#d0e8f0'
            },
            {
                'bg_color': '#fff8f0',  # 浅橙色
                'border_color': '#e0c0b0',
                'header_bg': '#f0e0d0'
            },
            {
                'bg_color': '#f8f8f0',  # 浅黄色
                'border_color': '#d0d0b0',
                'header_bg': '#e8e8d0'
            }
        ]
        # 从 options 中取出预先计算的行标记数据
        self.lines_with_markers = options.pop('lines_with_markers', [])
        # 添加新选项
        self.show_numbering = options.pop('show_numbering', True)
        self.show_dividers = options.pop('show_dividers', True)
        self.numbering_format = options.pop('numbering_format', '[1]')

    def _render_block(self, block_lines, block_num, total_blocks):
        """渲染一个完整的格式化文献块"""
        if not block_lines:
            return ""

        # 选择样式（循环使用）
        style = self.BLOCK_STYLES[block_num % len(self.BLOCK_STYLES)]
        
        # 生成编号
        number_html = ""
        if self.show_numbering:
            number_text = self.numbering_format.replace('1', str(block_num + 1))
            number_html = f'<span style="font-weight: bold; margin-right: 8px; color: #555;">{number_text}</span>'
        
        # 生成分割线
        divider = ""
        if self.show_dividers and block_num < total_blocks - 1:
            divider = '<div style="height: 1px; background: linear-gradient(90deg, transparent, #ccc, transparent); margin: 15px 0 5px 0;"></div>'
        
        # 文献标题 - 显示格式化后的状态
        block_header = (
            f'<div style="background: {style["header_bg"]}; padding: 10px 15px; '
            f'font-weight: bold; border-radius: 8px 8px 0 0; '
            f'border: 2px solid {style["border_color"]}; border-bottom: 1px dashed {style["border_color"]}; '
            f'color: #333; font-size: 12pt; display: flex; align-items: center;">'
            f'{number_html}📖 文献 {block_num + 1}'
            f'</div>'
        )

        # 文献内容
        block_content = []
        for line in block_lines:
            if line.strip():  # 只处理非空行
                line_html = (
                    f'<div style="padding: 8px 15px; '
                    f'border-left: 2px solid {style["border_color"]}; '
                    f'border-right: 2px solid {style["border_color"]}; '
                    f'font-family: \'Courier New\', monospace; '
                    f'white-space: pre-wrap;">'
                    f'{line}'
                    f'</div>'
                )
                block_content.append(line_html)

        # 包装整个文献块
        return (
            f'<div style="margin: 20px 0; border-radius: 8px;">'
            f'{block_header}'
            f'<div style="background-color: {style["bg_color"]}; '
            f'border: 2px solid {style["border_color"]}; border-top: none; '
            f'border-radius: 0 0 8px 8px;">'
            f'{"".join(block_content)}'
            f'</div>'
            f'{divider}'
            f'</div>'
        )

    def format_unencoded(self, tokensource, outfile):
        """
        覆盖格式化方法，显示格式化后的文献分割。
        """
        # 如果 outfile 为 None，使用 StringIO 作为缓冲区
        if outfile is None:
            string_buffer = StringIO()
            real_outfile = string_buffer
            should_return_string = True
        else:
            real_outfile = outfile
            should_return_string = False

        # 开始包装div
        real_outfile.write(
            '<div style="font-family: Arial, sans-serif; font-size: 11pt; line-height: 1.5; padding: 10px;">')

        current_block_lines = []
        block_index = 0
        total_blocks = sum(1 for _ in self.lines_with_markers if _[1])  # 计算总块数

        # 处理所有行
        for i, (line, is_new_block_start) in enumerate(self.lines_with_markers):
            if is_new_block_start and current_block_lines:
                # 渲染当前块
                real_outfile.write(self._render_block(current_block_lines, block_index, total_blocks))
                # 开始新块
                current_block_lines = [line]
                block_index += 1
            else:
                current_block_lines.append(line)

        # 渲染最后一个块
        if current_block_lines:
            real_outfile.write(self._render_block(current_block_lines, block_index, total_blocks))

        # 显示统计信息
        if block_index >= 0:
            real_outfile.write(
                f'<div style="text-align: center; color: #666; padding: 15px; '
                f'background: #f0f0f0; border-radius: 5px; margin-top: 20px; '
                f'font-size: 10pt;">'
                f'✅ 共检测到 {block_index + 1} 篇格式化文献'
                f'</div>'
            )

        real_outfile.write('</div>')

        # 如果使用了缓冲区，返回缓冲区内容
        if should_return_string:
            return string_buffer.getvalue()


class ReferenceProcessor:
    """
    一个专门用于处理和格式化学术文献引用的类。
    实现了中英文混合字体格式化、智能自动分割预览和Word导出功能。
    """

    # 定义中英文所需的字体
    CHINESE_FONT = 'SimSun'
    ENGLISH_FONT = 'Times New Roman'

    # **[START]** 新增：文献格式映射表 (用于下拉框和格式化逻辑)
    # 请确保您所有的引用格式都定义在这里
    FORMAT_MAP = {
        'GB/T 7714 (Author-Year)': {'format': 'author-year', 'style': 'gb7714-2015-author-date'},
        'GB/T 7714 (Numeric)': {'format': 'numeric', 'style': 'gb7714-2015-numeric'},
        # 请根据您的实际需求，在此处添加或修改更多的格式
        'APA 7th Edition': {'format': 'author-year', 'style': 'apa-7th-edition'},
        'Chicago 17th Edition': {'format': 'author-year', 'style': 'chicago-fullnote-bibliography'},
    }

    # **[END]** 新增

    def __init__(self):
        """
        初始化处理器，编译用于检测序号的正则表达式。
        """
        # 正则表达式匹配 [1], (1), 1. 三种格式的序号
        self.numbering_pattern = re.compile(r'^\s*(\[\d+\]|\(\d+\)|\d+\.)\s*')
        # 匹配 CJK 统一汉字，用于判断语言
        self.cjk_pattern = re.compile(r'[\u4e00-\u9fff]')
        print("DEBUG: ReferenceProcessor initialized.")

    def process_text(self, raw_text: str, format_name: str, numbering_format: str = "1.") -> tuple[str, str, bool]:
        """
        处理用户输入的完整原始文本。
        新增 numbering_format 参数支持 [1] 格式
        """
        print(f"DEBUG: Starting to process raw text with format: {numbering_format}")

        if not raw_text.strip():
            print("DEBUG: Raw text is empty.")
            return "", "", False

        lines = raw_text.strip().split('\n')

        styled_html_references = []
        plain_text_references = []
        was_stripped = False

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 1. 剥离已存在的序号
            cleaned_line, stripped, _ = self._strip_numbering(line)
            if stripped:
                was_stripped = True

            # 2. 对文献内容进行字符和符号规范化
            normalized_line = self._normalize_characters(cleaned_line)

            # 3. 生成两种输出
            plain_text_references.append(html.escape(normalized_line))
            styled_html_content = self._apply_mixed_font_styles(normalized_line)
            styled_html_references.append(styled_html_content)

        print(f"DEBUG: Stripped existing numbering: {was_stripped}")

        # 生成两种格式的文本 (传入 numbering_format)
        html_output = self._generate_html_list(styled_html_references)
        # 更新：传入编号格式用于界面预览
        plain_text_output = self._generate_plain_text_list(plain_text_references, numbering_format)

        return html_output, plain_text_output, was_stripped

    def _generate_plain_text_list(self, references: list, numbering_format: str = "1.") -> str:
        """
        生成用于在程序界面中预览的纯文本列表，支持自定义编号格式。
        """
        plain_list = []
        for i, ref in enumerate(references, 1):
            # 动态生成前缀：将格式中的 '1' 替换为当前序号
            # 例如: "[1]" -> "[2]", "(1)" -> "(2)"
            if "1" in numbering_format:
                prefix = numbering_format.replace("1", str(i))
            else:
                # 以此为兜底
                prefix = f"{i}."

            plain_list.append(f"{prefix} {html.unescape(ref)}")

        return "\n".join(plain_list)

    def _create_custom_numbering_xml(self, doc, format_string="[%1]"):
        """
        使用底层 XML 在 Word 文档中注入自定义的自动编号样式（如 [1]）。
        返回生成的 numId。
        """
        # 1. 确保文档已经初始化了 numbering 部分
        # 技巧：临时创建一个列表段落然后删除，强制 python-docx 初始化 numbering_part
        try:
            doc.part.numbering_part
        except Exception:
            p = doc.add_paragraph(style='List Number')
            doc._body._element.remove(p._element)

        numbering_part = doc.part.numbering_part

        # 获取底层的 numbering 元素
        # 注意：这里使用非公开 API (_numbering)，这是操作底层 XML 的唯一方式
        numbering_element = numbering_part.numbering_definitions._numbering

        # 2. 计算下一个可用的 ID
        # 获取当前所有 abstractNumId 的最大值
        abstract_num_ids = [int(e.get(qn('w:abstractNumId'))) for e in numbering_element.xpath('w:abstractNum')]
        next_abstract_id = max(abstract_num_ids) + 1 if abstract_num_ids else 1

        # 获取当前所有 numId 的最大值
        num_ids = [int(e.get(qn('w:numId'))) for e in numbering_element.xpath('w:num')]
        next_num_id = max(num_ids) + 1 if num_ids else 1

        # 3. 创建 abstractNum (定义样式外观)
        abstract_num = OxmlElement('w:abstractNum')
        abstract_num.set(qn('w:abstractNumId'), str(next_abstract_id))

        # 定义 Level 0 的样式
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), '0')

        # 起始编号 1
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        # 编号格式 decimal (1, 2, 3...)
        num_fmt = OxmlElement('w:numFmt')
        num_fmt.set(qn('w:val'), 'decimal')
        lvl.append(num_fmt)

        # 关键：定义显示文本，例如 [%1]
        lvl_text = OxmlElement('w:lvlText')
        lvl_text.set(qn('w:val'), format_string)
        lvl.append(lvl_text)

        # 对齐方式
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'left')
        lvl.append(jc)

        # 缩进设置 (悬挂缩进)
        p_pr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '420')  # 左缩进 (约0.74cm)
        ind.set(qn('w:hanging'), '420')  # 悬挂缩进 (对齐编号)
        p_pr.append(ind)
        lvl.append(p_pr)

        abstract_num.append(lvl)
        numbering_element.append(abstract_num)

        # 4. 创建 num (将样式实例化)
        num = OxmlElement('w:num')
        num.set(qn('w:numId'), str(next_num_id))

        abstract_num_ref = OxmlElement('w:abstractNumId')
        abstract_num_ref.set(qn('w:val'), str(next_abstract_id))
        num.append(abstract_num_ref)

        numbering_element.append(num)

        return next_num_id

    def export_to_word_file_with_custom_font(self, html_content: str, file_path: str,
                                             format_config: dict, numbering_format: str,
                                             english_font: str, english_size: float,
                                             chinese_font: str, chinese_size: float) -> bool:
        """
        使用自定义字体设置导出Word文件。
        支持多种自动编号格式，并修复了 XML 类型错误。
        """
        try:
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx 或 htmldocx 库未安装")

            if os.path.exists(file_path):
                try:
                    with open(file_path, 'a') as f:
                        pass
                except IOError as e:
                    raise PermissionError(f"文件可能被占用: {e}")

            doc = Document()
            self._setup_custom_font_styles(doc, format_config, english_font, english_size, chinese_font, chinese_size)

            # 1. 标题设置 (强制黑色)
            title = doc.add_heading('参考文献', 0)
            try:
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title.runs[0]
                title_run.font.size = Pt(format_config.get("title_font_size", 16))
                title_run.font.name = english_font
                title_run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
                title_run.font.color.rgb = RGBColor(0, 0, 0)
            except Exception as e:
                print(f"WARNING: Error formatting title: {e}")

            doc.add_paragraph().paragraph_format.space_after = Pt(format_config.get("title_margin_bottom", 20))

            references = self._extract_references_from_html(html_content)

            # 2. 准备自动编号样式
            # 定义界面选项到 Word XML 格式的映射 (%1 代表数字占位符)
            format_mapping = {
                "[1]": "[%1]",
                "(1)": "(%1)",
                "1)": "%1)",
                "<1>": "<%1>",
                "{1}": "{%1}"
            }

            custom_num_id = None

            # 只有当格式不是默认的 "1." 时，才进行 XML 注入
            if numbering_format != "1." and numbering_format in format_mapping:
                try:
                    xml_format_string = format_mapping[numbering_format]
                    custom_num_id = self._create_custom_numbering_xml(doc, xml_format_string)
                    print(f"DEBUG: Created custom numbering '{numbering_format}' with ID: {custom_num_id}")
                except Exception as e:
                    print(f"ERROR: Failed to create custom numbering XML: {e}")

            # 3. 循环写入文献
            for i, reference in enumerate(references, 1):
                if not reference or not reference.strip():
                    continue

                # 创建段落
                paragraph = doc.add_paragraph()

                # --- 应用编号 ---
                # 如果是自定义格式且注入成功
                if custom_num_id is not None:
                    # 使用自定义 XML 样式
                    p_pr = paragraph._p.get_or_add_pPr()
                    num_pr = p_pr.get_or_add_numPr()

                    # 设置 numId
                    num_id_elem = num_pr.get_or_add_numId()
                    # 🟢 【修复】必须传 int 类型，不能传 str
                    num_id_elem.val = int(custom_num_id)

                    # 设置 ilvl
                    ilvl_elem = num_pr.get_or_add_ilvl()
                    # 🟢 【修复】必须传 int 类型
                    ilvl_elem.val = 0

                else:
                    # 默认 "1." 样式 (或者注入失败回退)
                    paragraph.style = 'List Number'

                # 填充文献内容
                self._improved_mixed_font_text_custom(paragraph, reference,
                                                      english_font, english_size,
                                                      chinese_font, chinese_size)

                # 通用段落设置
                paragraph_format = paragraph.paragraph_format
                paragraph_format.line_spacing = format_config.get("line_spacing", 1.5)
                paragraph_format.space_after = Pt(format_config.get("item_spacing", 6))

            # 保存文件
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    doc.save(file_path)
                    return True
                except PermissionError as e:
                    if attempt < max_retries - 1:
                        import time
                        time.sleep(1)
                    else:
                        raise e
            return True

        except Exception as e:
            print(f"ERROR: Failed to export Word file: {str(e)}")
            import traceback
            traceback.print_exc()
            return False

    def _strip_numbering(self, line: str) -> tuple[str, bool, int]:
        """
        检查并从单个文献行中剥离已存在的序号。

        Args:
            line (str): 原始文献行。

        Returns:
            tuple[str, bool, int]: (清洗后的行, 是否被剥离, 剥离出的序号[整数])
        """
        match = self.numbering_pattern.match(line)
        if match:
            # 提取匹配到的序号字符串，并尝试转换为整数
            number_str = match.group(1).strip('[]() .')

            extracted_number = 0
            try:
                extracted_number = int(number_str)
            except ValueError:
                pass

            return line[match.end():].strip(), True, extracted_number

        return line.strip(), False, 0

    def _is_chinese_line(self, line: str, threshold: float = 0.1) -> bool:
        """
        简单的语言识别：如果文本中汉字字符数超过总有效字符数的阈值，则认为是中文行。
        """
        effective_line = re.sub(r'[\s\W_]', '', line)
        if not effective_line:
            return False

        chinese_chars = len(self.cjk_pattern.findall(line))
        return chinese_chars / len(effective_line) > threshold

    def _normalize_characters(self, line: str) -> str:
        """
        对文献条目进行字符和符号规范化，确保英文使用半角标点，中文保留全角标点。
        """
        normalized_line = line
        # 1. 通用全角转半角 (数字、括号等)
        for full_char, half_char in CHAR_MAPPING.items():
            normalized_line = normalized_line.replace(full_char, half_char)

        parts = []
        # 2. 分割：中文 | 非中文
        pattern = re.compile(r'([\u4e00-\u9fff]+)|([^\u4e00-\u9fff]+)')
        for match in pattern.finditer(normalized_line):
            chinese_part = match.group(1)
            other_part = match.group(2)

            if chinese_part:
                # 对中文部分进行标点全角规范化
                parts.append(self._normalize_chinese_punctuation(chinese_part))
            elif other_part:
                # 对非中文部分（英文、数字、标点）进行标点半角规范化
                parts.append(self._normalize_english_punctuation(other_part))

        normalized_line = "".join(parts)
        # 3. 替换多余的空格为一个空格，并去除首尾空格
        normalized_line = re.sub(r'\s+', ' ', normalized_line).strip()
        return normalized_line

    def _apply_mixed_font_styles(self, line: str) -> str:
        """
        遍历一行文本，使用 <span> 标签为中英文/数字片段分别应用不同字体。
        保留加粗和斜体格式，特别处理冒号后的文本。
        """
        # 首先处理加粗和斜体标记
        # 将 **text** 转换为 <b>text</b>
        line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
        # 将 *text* 转换为 <i>text</i>
        line = re.sub(r'\*(.*?)\*', r'<i>\1</i>', line)
        
        # 处理冒号后的文本，添加斜体
        parts = []
        segments = line.split(':')
        
        for i, segment in enumerate(segments):
            if i > 0:  # 在第一个冒号后的内容前后添加斜体标签
                parts.append(':<i>')
                parts.append(segment)
                parts.append('</i>')
            else:
                parts.append(segment)
        
        line = ''.join(parts)
        
        # 处理中英文混合
        result = []
        # 使用更精确的正则表达式匹配HTML标签
        pattern = re.compile(r'(<[^>]+>)|([\u4e00-\u9fff]+)|([a-zA-Z0-9\s]+)|([^<\u4e00-\u9fff\s]+)')
        
        for match in pattern.finditer(line):
            html_tag = match.group(1)
            chinese_part = match.group(2)
            english_part = match.group(3)
            other_part = match.group(4)
            
            if html_tag:
                # 直接添加HTML标签
                result.append(html_tag)
            elif chinese_part:
                result.append(
                    f'<span style="font-family: {self.CHINESE_FONT}; '
                    f'mso-hansi-font-family: {self.CHINESE_FONT}; '
                    f'mso-bidi-font-family: {self.CHINESE_FONT}; '
                    f'mso-ascii-font-family: {self.CHINESE_FONT};">'
                    f'{chinese_part}</span>'
                )
            elif english_part:
                result.append(
                    f'<span style="font-family: {self.ENGLISH_FONT}; '
                    f'mso-hansi-font-family: {self.ENGLISH_FONT}; '
                    f'mso-bidi-font-family: {self.ENGLISH_FONT}; '
                    f'mso-ascii-font-family: {self.ENGLISH_FONT};">'
                    f'{english_part}</span>'
                )
            elif other_part:
                # 处理标点符号
                is_chinese_punct = any(c in '，。：；！？' for c in other_part)
                font = self.CHINESE_FONT if is_chinese_punct else self.ENGLISH_FONT
                result.append(
                    f'<span style="font-family: {font}; '
                    f'mso-hansi-font-family: {font}; '
                    f'mso-bidi-font-family: {font}; '
                    f'mso-ascii-font-family: {font};">'
                    f'{other_part}</span>'
                )
        
        return ''.join(result)

    def _generate_html_list(self, styled_html_references: list[str]) -> str:
        """
        生成 Word 可识别的 HTML 自动编号列表，使用Word原生自动编号。
        """
        if not styled_html_references:
            return ""

        # 构建列表项 - 使用Word自动编号
        list_items = []
        for styled_html_content in styled_html_references:
            list_item = f'''
            <li style="margin-bottom: 6pt;">
                {styled_html_content}
            </li>
            '''
            list_items.append(list_item)

        # 使用Word原生自动编号的CSS
        style = f"""
        <style>
            /* Word原生自动编号样式 */
            ol {{
                list-style-type: none;
                counter-reset: item;
                margin: 0;
                padding: 0;
            }}
            li {{
                display: block;
                margin-bottom: 6pt;
                margin-left: 0;
                padding-left: 24pt;
                text-indent: -12pt;
            }}
            li:before {{
                content: counter(item) ". ";
                counter-increment: item;
                display: inline-block;
                width: 12pt;
                margin-left: -12pt;
                font-family: '{self.ENGLISH_FONT}';
                font-weight: normal;
            }}
            /* 确保Word兼容性 */
            @list l0:level1 {{
                mso-level-number-format: decimal;
                mso-level-text: "%1.";
                mso-level-tab-stop: 36.0pt;
                mso-level-number-position: left;
                margin-left: 36.0pt;
                text-indent: -18.0pt;
            }}
        </style>
        """

        # 最终的 Word HTML 结构 - 使用Word原生列表
        html_string = f"""
        <html xmlns:o="urn:schemas-microsoft-com:office:office"
              xmlns:w="urn:schemas-microsoft-com:office:word"
              xmlns:m="http://schemas.microsoft.com/office/2004/12/omml">
        <head>
            <meta charset="UTF-8">
            {style}
        </head>
        <body>
        <!--[if supportLists]>
        <ol style="list-style-type: decimal;">
        <![endif]-->
        <!--[if supportLists]><!-->
        {'<!--[endif]--><!--[if supportLists]-->'.join(list_items)}
        <!--[endif]-->
        <!--[if supportLists]>
        </ol>
        <![endif]-->
        </body>
        </html>
        """
        return html_string



    def get_formatted_split_preview(self, raw_text: str, format_name: str) -> str:
        """
        先格式化文献，再进行分割预览。

        Args:
            raw_text (str): 原始输入的文献文本
            format_name (str): 格式名称（保留参数，用于兼容性）

        Returns:
            str: 包含格式化后文献分割预览的 HTML 字符串
        """
        if not raw_text.strip():
            return ""

        # 1. 先进行完整的格式化处理
        html_output, plain_text_output, was_stripped = self.process_text(raw_text, format_name)

        # 2. 从纯文本输出中获取格式化后的文献列表
        formatted_lines = plain_text_output.split('\n')

        # 3. 对格式化后的文本进行分割检测
        lines_with_markers = self._detect_boundary_lines(formatted_lines)

        # 4. 使用自定义 Formatter 生成分割预览
        formatter = ReferenceBlockFormatter(lines_with_markers=lines_with_markers)
        lexer = ReferenceBlockLexer()

        # 重新组合文本用于 Pygments 处理
        formatted_text = '\n'.join([line for line, _ in lines_with_markers])

        html_output = html.unescape(formatter.format(lexer.get_tokens_unprocessed(0, formatted_text), outfile=None))

        print("DEBUG: Formatted split preview generation finished.")
        return html_output

    def _detect_boundary_lines(self, lines: list[str]) -> list[tuple[str, bool]]:
        """
        改进的边界检测：基于格式化后的编号识别不同的文献。
        """
        lines_with_markers = []

        if not lines:
            return lines_with_markers

        # 强制第一行总是新块的开始
        lines_with_markers.append((lines[0], True))

        for i in range(1, len(lines)):
            current_line = lines[i]
            current_stripped = current_line.strip()

            is_new_block = False

            # 主要分割条件：检测到格式化后的编号模式
            if current_stripped:
                # 检查是否包含格式化后的文献编号
                if self._has_formatted_numbering(current_stripped):
                    is_new_block = True
                # 检查空行后的内容（可能是新文献）
                elif i > 0 and not lines[i - 1].strip() and current_stripped:
                    is_new_block = True

            # 保持原始行内容，并标记是否为新块开始
            lines_with_markers.append((current_line, is_new_block))

        return lines_with_markers

    def _has_formatted_numbering(self, line: str) -> bool:
        """
        检测行是否包含格式化后的编号模式。
        """
        # 匹配格式化后的常见编号模式
        formatted_patterns = [
            r'^\[\d+\]',  # [1]
            r'^\d+\.',  # 1.
            r'^\(\d+\)',  # (1)
        ]

        for pattern in formatted_patterns:
            if re.match(pattern, line.strip()):
                return True

        return False

    def _normalize_english_punctuation(self, text: str) -> str:
        """
        英文标点规范化：确保使用半角标点
        """
        full_to_half = {
            '，': ',', '。': '.', '：': ':', '；': ';',
            '？': '?', '！': '!', '（': '(', '）': ')',
            '【': '[', '】': ']', '―': '-', '–': '-',
            '—': '-', '．': '.', '·': '.', '∶': ':',
            '＠': '@', '＃': '#', '＄': '$', '％': '%',
            '＾': '^', '＆': '&', '＊': '*', '＋': '+',
            '＝': '=', '～': '~', '　': ' '
        }

        full_digits = '０１２３４５６７８９'
        half_digits = '0123456789'
        digit_mapping = str.maketrans(full_digits, half_digits)

        result = text.translate(digit_mapping)
        for full, half in full_to_half.items():
            result = result.replace(full, half)

        return result

    def _normalize_chinese_punctuation(self, text: str) -> str:
        """
        中文标点规范化：确保使用全角标点
        """
        half_to_full = {
            ',': '，', '.': '。', ':': '：', ';': '；',
            '?': '？', '!': '！', '(': '（', ')': '）',
            '[': '【', ']': '】'
        }

        result = text
        for half, full in half_to_full.items():
            result = result.replace(half, full)

        return result



    def _setup_custom_font_styles(self, doc, format_config, english_font, english_size, chinese_font, chinese_size):
        """设置自定义字体样式"""
        try:
            # 设置默认样式
            style = doc.styles['Normal']
            font = style.font
            font.name = english_font
            font.size = Pt(english_size)

            # 设置东亚字体（中文字体）
            font._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)

            # 设置列表样式
            try:
                # 尝试获取或创建列表样式
                list_style = doc.styles.add_style('ReferenceList', WD_STYLE_TYPE.PARAGRAPH)
                list_style.base_style = doc.styles['List Number']
                list_font = list_style.font
                list_font.name = english_font
                list_font.size = Pt(english_size)
                list_font._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
            except Exception as e:
                print(f"WARNING: Error setting list style: {e}")

            # 设置页面边距
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
        except Exception as e:
            print(f"WARNING: Error in font setup: {e}")

    def _improved_mixed_font_text_custom(self, paragraph, text, english_font, english_size, chinese_font, chinese_size):
        """
        使用自定义字体的改进中英文混合字体设置方法。
        """
        try:
            chinese_chars = []
            english_chars = []
            current_type = None

            for char in text:
                if self._is_chinese_char(char):
                    char_type = 'chinese'
                elif char.isalpha() or char.isdigit() or char in ' .,;:!?\'"()-':
                    char_type = 'english'
                else:
                    char_type = 'chinese'

                if current_type != char_type:
                    if chinese_chars or english_chars:
                        self._add_text_segment_custom(paragraph, chinese_chars, english_chars,
                                                      english_font, english_size, chinese_font, chinese_size)
                        chinese_chars = []
                        english_chars = []
                    current_type = char_type

                if char_type == 'chinese':
                    chinese_chars.append(char)
                else:
                    english_chars.append(char)

            if chinese_chars or english_chars:
                self._add_text_segment_custom(paragraph, chinese_chars, english_chars,
                                              english_font, english_size, chinese_font, chinese_size)
        except Exception as e:
            print(f"WARNING: Error in mixed font text: {e}")
            # 如果出错，直接添加文本
            paragraph.add_run(text)

    def _add_text_segment_custom(self, paragraph, chinese_chars, english_chars,
                                 english_font, english_size, chinese_font, chinese_size):
        """使用自定义字体添加文本片段"""
        try:
            if chinese_chars:
                chinese_text = ''.join(chinese_chars)
                run = paragraph.add_run(chinese_text)
                run.font.name = chinese_font
                run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
                run.font.size = Pt(chinese_size)

            if english_chars:
                english_text = ''.join(english_chars)
                run = paragraph.add_run(english_text)
                run.font.name = english_font
                run.font.size = Pt(english_size)
        except Exception as e:
            print(f"WARNING: Error adding text segment: {e}")

    def _is_chinese_char(self, char):
        """判断字符是否为中文字符"""
        # 基本CJK统一汉字
        if '\u4e00' <= char <= '\u9fff':
            return True
        # CJK扩展A
        if '\u3400' <= char <= '\u4dbf':
            return True
        # 中文标点符号
        if char in '，。！？；：「」『』【】（）《》':
            return True
        return False

    def _extract_references_from_html(self, html_content):
        """
        从HTML内容中提取参考文献列表。
        """
        try:
            # 提取body内容
            body_content = self._extract_body_content(html_content)

            # 使用正则表达式提取列表项内容
            references = []
            li_pattern = re.compile(r'<li[^>]*>(.*?)</li>', re.DOTALL | re.IGNORECASE)

            for match in li_pattern.finditer(body_content):
                li_content = match.group(1)

                # 移除HTML标签但保留文本内容
                clean_text = re.sub('<[^<]+?>', '', li_content)
                # 解码HTML实体
                clean_text = html.unescape(clean_text)
                # 移除多余的空白字符
                clean_text = re.sub(r'\s+', ' ', clean_text).strip()

                if clean_text:
                    references.append(clean_text)

            print(f"DEBUG: Extracted {len(references)} references from HTML")
            return references

        except Exception as e:
            print(f"WARNING: Error extracting references from HTML: {e}")
            return ["文献内容提取失败"]

    def _extract_body_content(self, html_content: str) -> str:
        """
        从完整的HTML内容中提取body部分。
        """
        try:
            # 使用正则表达式提取body内容
            body_match = re.search(r'<body[^>]*>(.*?)</body>', html_content, re.DOTALL | re.IGNORECASE)
            if body_match:
                return body_match.group(1).strip()
            else:
                # 如果没有body标签，返回原始内容
                return html_content
        except Exception as e:
            print(f"WARNING: Error extracting body content: {e}")
            return html_content